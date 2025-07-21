"""Word document loader service."""

import asyncio
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_community.document_loaders import UnstructuredWordDocumentLoader
from langchain_core.documents import Document

from .base_loader import BaseDocumentLoader

logger = logging.getLogger(__name__)


class WordDocumentLoader(BaseDocumentLoader):
    """Document loader for Word files (.doc/.docx)."""

    def __init__(self, encoding_name: str = "cl100k_base"):
        """Initialize the Word document loader.

        Args:
            encoding_name: The tokenizer encoding to use for token counting
        """
        super().__init__(encoding_name)
        self.logger = logging.getLogger(__name__)

    async def load_document(
        self,
        file_path: str,
        password: Optional[str] = None,
        preserve_layout: bool = True,
        encoding_info: Optional[Dict[str, Any]] = None,
    ) -> List[Document]:
        """Load a Word document.

        Args:
            file_path: Path to the Word file
            password: Not currently supported for Word files
            preserve_layout: Whether to preserve Word document layout information
            encoding_info: Not used for Word files

        Returns:
            List of Document objects
        """
        self.logger.info(f"Loading Word document: {file_path}")

        if preserve_layout:
            return await self._load_word_with_elements_mode(file_path)
        else:
            return await self._load_word_standard_mode(file_path)

    async def _load_word_with_elements_mode(self, file_path: str) -> List[Document]:
        """Load Word document with element-based processing for layout preservation."""
        try:
            # Try using python-docx for better structure extraction
            documents = await self._load_word_with_python_docx(file_path)
            if documents:
                return documents
        except Exception as e:
            self.logger.warning(
                f"python-docx loading failed: {str(e)}, falling back to unstructured"
            )

        # Fallback to unstructured loader with elements mode
        try:
            loader = UnstructuredWordDocumentLoader(file_path, mode="elements")
            elements = loader.load()

            # Group elements into logical sections
            sections = self._create_section_documents(elements, file_path)

            self.logger.info(
                f"Successfully loaded Word document with elements mode: {len(sections)} sections"
            )
            return sections

        except Exception as e:
            self.logger.error(f"Error loading Word document with elements mode: {str(e)}")
            # Final fallback to standard mode
            return await self._load_word_standard_mode(file_path)

    def _create_section_documents(self, elements: List[Document], file_path: str) -> List[Document]:
        """Create section-based documents from unstructured elements."""
        sections = []
        current_section = {
            "content": [],
            "metadata": {
                "source": file_path,
                "file_type": "word",
                "layout_preserved": True,
                "formatting": {"headings": [], "lists": [], "tables": [], "text_blocks": []},
            },
        }

        section_index = 0

        for element in elements:
            element_type = element.metadata.get("category", "text")
            content = element.page_content.strip()

            if not content:
                continue

            # Start new section on major headings
            if element_type == "Title" or (element_type == "Header" and current_section["content"]):
                if current_section["content"]:
                    sections.append(self._finalize_section_document(current_section, section_index))
                    section_index += 1

                current_section = {
                    "content": [],
                    "metadata": {
                        "source": file_path,
                        "file_type": "word",
                        "layout_preserved": True,
                        "formatting": {
                            "headings": [],
                            "lists": [],
                            "tables": [],
                            "text_blocks": [],
                        },
                    },
                }

            # Add content to current section
            current_section["content"].append(content)

            # Track formatting elements
            formatting = current_section["metadata"]["formatting"]
            if element_type in ["Title", "Header"]:
                formatting["headings"].append(
                    {
                        "text": content,
                        "level": 1 if element_type == "Title" else 2,
                        "position": len(current_section["content"]) - 1,
                    }
                )
            elif element_type == "ListItem":
                formatting["lists"].append(
                    {
                        "text": content,
                        "type": "bulleted",
                        "position": len(current_section["content"]) - 1,
                    }
                )
            elif element_type == "Table":
                table_data = self._process_table_element(element, len(formatting["tables"]))
                formatting["tables"].append(table_data)
            else:
                formatting["text_blocks"].append(
                    {
                        "text": content,
                        "type": element_type.lower(),
                        "position": len(current_section["content"]) - 1,
                    }
                )

        # Add final section
        if current_section["content"]:
            sections.append(self._finalize_section_document(current_section, section_index))

        return sections

    def _finalize_section_document(self, section: Dict, section_index: int) -> Document:
        """Finalize a section into a Document object."""
        content = "\n".join(section["content"])
        metadata = section["metadata"].copy()
        metadata.update(
            {
                "section_index": section_index,
                "word_count": len(content.split()),
                "char_count": len(content),
            }
        )

        return Document(page_content=content, metadata=metadata)

    def _process_table_element(self, element: Document, table_index: int) -> Dict[str, Any]:
        """Process a table element and extract structured data."""
        content = element.page_content

        # Simple table parsing - split by lines and assume tab/space separation
        lines = [line.strip() for line in content.split("\n") if line.strip()]
        table_data = []

        for line in lines:
            # Try different separators
            if "\t" in line:
                row = [cell.strip() for cell in line.split("\t")]
            else:
                # Fallback to multiple spaces
                row = [cell.strip() for cell in line.split("  ") if cell.strip()]

            if row:
                table_data.append(row)

        return {
            "index": table_index,
            "rows": len(table_data),
            "columns": len(table_data[0]) if table_data else 0,
            "data": table_data,
            "raw_content": content,
        }

    async def _load_word_standard_mode(self, file_path: str) -> List[Document]:
        """Load Word document using standard unstructured loader."""
        try:
            loader = UnstructuredWordDocumentLoader(file_path)
            documents = loader.load()

            # Add basic metadata
            for doc in documents:
                doc.metadata.update(
                    {
                        "source": file_path,
                        "file_type": "word",
                        "layout_preserved": False,
                    }
                )

            self.logger.info(
                f"Successfully loaded Word document with standard mode: {len(documents)} documents"
            )
            return documents

        except Exception as e:
            self.logger.error(f"Error loading Word document: {str(e)}")
            raise ValueError(f"Failed to load Word document: {str(e)}")

    async def _load_word_with_python_docx(self, file_path: str) -> List[Document]:
        """Load Word document using python-docx for better structure extraction."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            self.logger.warning("python-docx not available")
            return []

        try:
            doc = DocxDocument(file_path)

            # Extract structured content
            structured_elements = await self._extract_structured_word_content(doc)

            # Group into sections
            sections = group_word_content_into_sections(structured_elements)

            # Convert to Document objects
            documents = []
            for i, section in enumerate(sections):
                metadata = {
                    "source": file_path,
                    "section_index": i,
                    "file_type": "word",
                    "layout_preserved": True,
                    "formatting": section["formatting"],
                    "word_count": section["word_count"],
                    "char_count": section["char_count"],
                }

                documents.append(Document(page_content=section["content"], metadata=metadata))

            # Extract document metadata
            doc_metadata = await self._extract_word_metadata(file_path)
            for document in documents:
                document.metadata["document_metadata"] = doc_metadata

            self.logger.info(
                f"Successfully loaded Word document with python-docx: {len(documents)} sections"
            )
            return documents

        except Exception as e:
            self.logger.error(f"Error loading Word document with python-docx: {str(e)}")
            return []

    async def _extract_structured_word_content(self, doc) -> List[Dict[str, Any]]:
        """Extract structured content from python-docx Document."""
        elements = []

        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue

            element_type = determine_word_element_type(paragraph)
            formatting = extract_paragraph_formatting(paragraph)

            elements.append(
                {
                    "type": element_type,
                    "content": paragraph.text.strip(),
                    "formatting": formatting,
                }
            )

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_content = extract_table_content(table)
            elements.append(
                {
                    "type": "table",
                    "content": f"[TABLE_{table_idx}]",
                    "table_data": table_content,
                    "formatting": {},
                }
            )

        return elements

    async def _extract_word_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from Word document."""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            core_props = doc.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "file_size": Path(file_path).stat().st_size,
            }

            return metadata

        except Exception as e:
            self.logger.error(f"Error extracting Word metadata: {str(e)}")
            return {"error": str(e)}

    async def extract_word_with_batch_processing(
        self, file_paths: List[str], max_concurrent: int = 3
    ) -> List[List[Document]]:
        """Process multiple Word documents concurrently.

        Args:
            file_paths: List of paths to Word files
            max_concurrent: Maximum number of concurrent processing tasks

        Returns:
            List of document lists, one for each input file
        """
        semaphore = asyncio.Semaphore(max_concurrent)

        async def process_single_file(file_path: str) -> List[Document]:
            try:
                return await self.load_document(file_path)
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {str(e)}")
                return []

        async def process_with_semaphore(file_path: str) -> List[Document]:
            async with semaphore:
                return await process_single_file(file_path)

        tasks = [process_with_semaphore(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Handle exceptions
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                self.logger.error(f"Failed to process {file_paths[i]}: {str(result)}")
                processed_results.append([])
            else:
                processed_results.append(result)

        return processed_results

    def get_word_processing_stats(self) -> Dict[str, Any]:
        """Get statistics about Word processing capabilities.

        Returns:
            Dictionary containing processing statistics and capabilities
        """
        stats = {
            "supported_extensions": self.get_supported_extensions(),
            "supported_mime_types": self.get_supported_mime_types(),
            "layout_preservation": True,
            "batch_processing": True,
            "metadata_extraction": True,
        }

        # Check for optional dependencies
        try:
            stats["python_docx_available"] = True
        except ImportError:
            stats["python_docx_available"] = False

        return stats

    def get_supported_mime_types(self) -> List[str]:
        """Get the MIME types supported by this loader.

        Returns:
            List of supported MIME types
        """
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def get_supported_extensions(self) -> List[str]:
        """Get the file extensions supported by this loader.

        Returns:
            List of supported file extensions
        """
        return [".docx", ".doc"]
